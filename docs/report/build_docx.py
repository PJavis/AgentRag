#!/usr/bin/env python
"""Assemble chapter markdown files into a Vietnamese thesis-format .docx.

Format: A4, margins left 3cm / right 2cm / top 2cm / bottom 2cm,
Times New Roman 13pt, 1.5 line spacing, justified body, page numbers,
auto TOC field (press F9 / "Update field" in Word to populate).
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE = Path(__file__).parent
CHAPTERS = ["ch1.md", "ch2.md", "ch3.md", "ch4.md", "ch5.md", "ch6.md"]
REFS = "refs.md"
OUT = Path(sys.argv[1]) if len(sys.argv) > 1 else BASE / "BaoCao_AgentRag.docx"

FONT = "Times New Roman"


def set_font(run, size=13, bold=False, italic=False, mono=False, color=None):
    run.font.name = "Consolas" if mono else FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    fname = "Consolas" if mono else FONT
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), fname)


def body_par(doc, justify=True, spacing=1.5, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.space_after = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


INLINE_RE = re.compile(
    r"(\*\*\*(?P<bi>.+?)\*\*\*|\*\*(?P<b>.+?)\*\*|\*(?P<i>[^*]+?)\*|`(?P<c>[^`]+?)`|\$(?P<m>[^$]+?)\$)"
)


def add_inline(p, text, size=13, base_bold=False):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            set_font(p.add_run(text[pos:m.start()]), size=size, bold=base_bold)
        if m.group("bi") is not None:
            set_font(p.add_run(m.group("bi")), size=size, bold=True, italic=True)
        elif m.group("b") is not None:
            set_font(p.add_run(m.group("b")), size=size, bold=True)
        elif m.group("i") is not None:
            set_font(p.add_run(m.group("i")), size=size, italic=True, bold=base_bold)
        elif m.group("c") is not None:
            set_font(p.add_run(m.group("c")), size=12, mono=True, bold=base_bold)
        elif m.group("m") is not None:
            set_font(p.add_run(m.group("m")), size=size, italic=True, bold=base_bold)
        pos = m.end()
    if pos < len(text):
        set_font(p.add_run(text[pos:]), size=size, bold=base_bold)


def add_heading(doc, level, text, first_chapter=False):
    p = doc.add_paragraph()
    if level == 1:
        if not first_chapter:
            p.paragraph_format.page_break_before = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(18)
        set_font(p.add_run(text.upper()), size=16, bold=True)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        set_font(p.add_run(text), size=14, bold=True)
    elif level == 3:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(text), size=13, bold=True, italic=True)
    else:
        p.paragraph_format.space_before = Pt(6)
        set_font(p.add_run(text), size=13, italic=True)
    p.style = doc.styles[f"Heading {min(level, 4)}"]
    # style overrides direct formatting of style defaults; re-apply run font
    for run in p.runs:
        if level == 1:
            set_font(run, size=16, bold=True)
        elif level == 2:
            set_font(run, size=14, bold=True)
        elif level == 3:
            set_font(run, size=13, bold=True, italic=True)
        else:
            set_font(run, size=13, italic=True)
    return p


def add_table(doc, header_cells, rows):
    ncols = len(header_cells)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[j]
        cell.paragraphs[0].text = ""
        add_inline(cell.paragraphs[0], cell_text.strip(), size=12, base_bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows, start=1):
        for j in range(ncols):
            cell = table.rows[i].cells[j]
            cell.paragraphs[0].text = ""
            add_inline(cell.paragraphs[0], (row[j] if j < len(row) else "").strip(), size=12)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.75)
        set_font(p.add_run(line if line else " "), size=10.5, mono=True)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "F2F2F2")
        p._element.get_or_add_pPr().append(shd)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def split_table_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def render_markdown(doc, md_text, first_chapter=False):
    lines = md_text.splitlines()
    i, n = 0, len(lines)
    seen_h1 = False
    while i < n:
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("```"):
            block = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, block)
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            add_heading(doc, level, m.group(2).strip(),
                        first_chapter=(first_chapter and level == 1 and not seen_h1))
            if level == 1:
                seen_h1 = True
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|[\s:|-]+\|?$", lines[i + 1].strip()):
            header = split_table_row(stripped)
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i].strip()))
                i += 1
            add_table(doc, header, rows)
            continue
        m = re.match(r"^([-*+]|\d+[.)])\s+(.*)$", stripped)
        if m:
            p = body_par(doc, space_after=2)
            p.style = doc.styles["List Bullet" if m.group(1) in "-*+" else "List Number"]
            p.paragraph_format.line_spacing = 1.5
            add_inline(p, m.group(2))
            i += 1
            continue
        if stripped.startswith(">"):
            p = body_par(doc)
            p.paragraph_format.left_indent = Cm(1)
            add_inline(p, stripped.lstrip("> ").strip())
            i += 1
            continue
        # paragraph: merge consecutive plain lines
        buf = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith(("#", "|", "```", ">", "- ", "* ", "+ "))
                    or re.match(r"^\d+[.)]\s", nxt)):
                break
            buf.append(nxt)
            i += 1
        p = body_par(doc)
        p.paragraph_format.first_line_indent = Cm(1)
        add_inline(p, " ".join(buf))


def add_page_numbers(doc):
    footer = doc.sections[-1].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for el, attrs, text in (
        ("w:fldChar", {"w:fldCharType": "begin"}, None),
        ("w:instrText", {"xml:space": "preserve"}, " PAGE "),
        ("w:fldChar", {"w:fldCharType": "end"}, None),
    ):
        node = OxmlElement(el)
        for k, v in attrs.items():
            node.set(qn(k), v)
        if text:
            node.text = text
        run._element.append(node)
    set_font(run, size=12)


def add_toc_field(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    hint = OxmlElement("w:t")
    hint.text = "Nhấn chuột phải → Update Field (hoặc F9) để tạo mục lục tự động."
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    for node in (fld_begin, instr, fld_sep, hint, fld_end):
        run._element.append(node)
    set_font(run, size=13)


def title_page(doc):
    def center(text, size, bold=True, before=0, after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        set_font(p.add_run(text), size=size, bold=bold)
        return p

    center("TRƯỜNG ĐẠI HỌC BÁCH KHOA HÀ NỘI", 14, before=12)
    center("TRƯỜNG CÔNG NGHỆ THÔNG TIN VÀ TRUYỀN THÔNG", 13, after=48)
    center("BÁO CÁO ĐỒ ÁN", 20, before=60)
    center("XÂY DỰNG HỆ THỐNG HỎI ĐÁP TÀI LIỆU Y TẾ TIẾNG VIỆT", 16, before=24)
    center("DỰA TRÊN KIẾN TRÚC RETRIEVAL-AUGMENTED GENERATION", 16)
    center("VÀ AGENT ĐA BƯỚC (HỆ THỐNG AGENTRAG/VITAL)", 16, after=72)
    center("Sinh viên thực hiện: Nguyễn Quốc Dũng", 13, bold=False, before=48)
    center("Giảng viên hướng dẫn: (điền tên)", 13, bold=False)
    center("Hà Nội, 07/2026", 13, bold=False, before=48)
    doc.add_page_break()


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(13)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    title_page(doc)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(h.add_run("MỤC LỤC"), size=16, bold=True)
    add_toc_field(doc)
    doc.add_page_break()

    for idx, name in enumerate(CHAPTERS):
        path = BASE / name
        if not path.exists():
            print(f"!! missing {name}", file=sys.stderr)
            continue
        render_markdown(doc, path.read_text(encoding="utf-8"), first_chapter=(idx == 0))

    refs = BASE / REFS
    if refs.exists():
        add_heading(doc, 1, "TÀI LIỆU THAM KHẢO")
        text = refs.read_text(encoding="utf-8")
        text = re.sub(r"^#.*$", "", text, flags=re.MULTILINE)
        for line in text.splitlines():
            line = line.strip().lstrip("-* ").strip()
            if not line:
                continue
            p = body_par(doc, space_after=4)
            add_inline(p, line)

    add_page_numbers(doc)
    doc.save(OUT)

    words = 0
    for name in CHAPTERS:
        path = BASE / name
        if path.exists():
            words += len(path.read_text(encoding="utf-8").split())
    print(f"saved {OUT} — chapter words total: {words} (~{words // 450} pages of body text)")


if __name__ == "__main__":
    main()
